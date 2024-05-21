import os  #operating system library
from docx import Document  #Word documents library
from datetime import datetime
import re  #regular expressions library
import threading  #multi-threading

def create_dir(ip): # create main direcoty with ip, date and time
    current_datetime = datetime.now().strftime("%H-%M-%d_%m-%y")#Get the current date and time 
    dir_name = f"{current_datetime}_{ip}"     # Combine date and ip
    os.makedirs(dir_name)     # Create the directory
    return dir_name

def run_whatweb(ip, dir_name, result_holder): # Function to execute WhatWeb and save output
    output_file = os.path.join(dir_name, "whatweb.txt") # Define the path for the output file
    url = "http://" + ip  # concatenate and turn ip to url
    os.system(f"whatweb {url} > {output_file}")  # run command and save the output
    print(f"WhatWeb output was saved to {output_file}")
    result_holder['whatweb'] = output_file

# same comments from run_whatweb are applied here
def run_nmap(ip, dir_name, result_holder):
    output_file = os.path.join(dir_name, "nmap.txt")
    os.system(f"nmap -sV -O -p- {ip} > {output_file}")
    print(f"Nmap output was saved to {output_file}")
    result_holder['nmap'] = output_file

    # Call the function extract and write versions from the Nmap output
    with open(output_file, 'r') as f:
        nmap_output = f.read()
        extract_versions(dir_name, nmap_output)

# function to extract versions from Nmap output
def extract_versions(dir_name, nmap_output):
    # Regular expressions to extract Apache and PHP versions
    apache_regex = r"Apache httpd ([\d.]+)"
    php_regex = r"PHP/([\d.]+)"

    # wxtract apache and PHP versions using regular expressions
    apache_match = re.search(apache_regex, nmap_output)
    apache_version = apache_match.group(1) if apache_match else ""
    
    php_match = re.search(php_regex, nmap_output)
    php_version = php_match.group(1) if php_match else ""

    # write Apache and PHP versions to a file
    with open(os.path.join(dir_name, "versions.txt"), "w") as f:
        f.write(f"Apache version: {apache_version}\n")
        f.write(f"PHP version: {php_version}\n")
    print("Versions extracted to 'versions.txt'.")

    # run searchsploit for Apache and PHP versions
    run_searchsploit(dir_name, apache_version, php_version)

def run_searchsploit(dir_name, apache_version, php_version):
    # execute searchsploit for apache and PHP versions and save the results to a file
    os.system(f"searchsploit 'Apache {apache_version}' > {os.path.join(dir_name, 'sploits.txt')}")
    os.system(f"searchsploit 'PHP {php_version}' >> {os.path.join(dir_name, 'sploits.txt')}")
    sploits_file = os.path.join(dir_name, "sploits.txt")
    print("Searchsploit results were saved to 'sploits.txt'.")
    return '\n'.join(sploits_file)

# same comments from run_whatweb are applied here
def run_nikto(ip, dir_name, result_holder):
    output_file = os.path.join(dir_name, "nikto.txt")
    url = "http://" + ip 
    os.system(f"nikto -h {url} > {output_file}")
    print(f"Nikto output was saved to {output_file}")
    result_holder['nikto'] = output_file
    
# same comments from run_whatweb are applied here
def run_dirb(ip, dir_name, result_holder):
    output_file = os.path.join(dir_name, "dirb.txt")
    url = "http://" + ip
    os.system(f"dirb {url} > {output_file}")
    print(f"Dirb output was saved to {output_file}")
    result_holder['dirb'] = output_file

# function to parse Dirb output and filter relevant lines
def parse_dirb_output(output_file):
    try:
        # Open the Dirb output file
        with open(output_file, 'r') as f:
            # Read all lines from the file
            lines = f.readlines()
            # Filter lines containing the word 'directory' and remove whitespaces
            filtered_lines = [line.strip() for line in lines if 'directory' in line.lower()]
        # turn filtered lines into a single string separated by newline characters
        return '\n'.join(filtered_lines)
    except Exception as e:
        print(f"Error parsing Dirb output: {e}")
        return ''

#generate a Word document report with the results
def generate_word_document(ip, dir_name, whatweb_output, nmap_output, nikto_output, dirb_output):
    document = Document() # create a new Word document 
    document.add_heading(f"Web Application Testing Report - {ip}", level=1) # add a heading
    
    # loop through each tool output and add it to the document
    for name, output_file in [("WhatWeb Output", whatweb_output), ("Nmap Output", nmap_output), ("Nikto Output", nikto_output)]:
        document.add_heading(name, level=2)# add a subheading with the tool name
        with open(output_file, "r", encoding="utf-8", errors="ignore") as f: # open the output file 
            text = re.sub(r'[^\x20-\x7E]', '', f.read())  # read the file contents and remove non-ASCII characters
            document.add_paragraph(text) # add the cleaned text to the document as a paragraph

    document.add_heading("Dirb Output", level=2) #add a subheading with the tool name
    dirb_filtered_output = parse_dirb_output(dirb_output) # Parse and filter Dirb output
    document.add_paragraph(dirb_filtered_output) # Add the filtered Dirb output to the document as a paragraph
   
    document.add_heading("Searchsploit Results", level=2) # create sploits output file   
    sploits_file_path = os.path.join(dir_name, "sploits.txt")
    with open(sploits_file_path, "r", encoding="utf-8", errors="ignore") as f:
        sploits_text = re.sub(r'[^\x20-\x7E]', '', f.read()) # read the contents and remove non-ASCII characters
        document.add_paragraph(sploits_text) # add the cleaned text to the document as a paragraph
        
        
        # Define the path for the final Word document
        output_file = os.path.join(dir_name, f"{ip}_Security_Assessment_Report.docx")
        document.save(output_file) # Save the document 
    print(f"Word document saved as {output_file}")


def main():
    # ask for ip
    ip = input("Introcude the target IP: ")
    # Create a directory to store assessment results
    dir_name = create_dir(ip)
    
    # Create a dictionary to store the results of the tools
    result_holder = {}
    
    # Create threads to execute each tool simultaneously
    whatweb_thread = threading.Thread(target=run_whatweb, args=(ip, dir_name, result_holder))
    nmap_thread = threading.Thread(target=run_nmap, args=(ip, dir_name, result_holder))
    nikto_thread = threading.Thread(target=run_nikto, args=(ip, dir_name, result_holder))
    dirb_thread = threading.Thread(target=run_dirb, args=(ip, dir_name, result_holder))
    
    # Start the threads 
    whatweb_thread.start()
    nmap_thread.start()
    nikto_thread.start()
    dirb_thread.start()
    
    # Wait for all threads to finish
    whatweb_thread.join()
    nmap_thread.join()
    nikto_thread.join()
    dirb_thread.join()
    
    # Extract the results from the dictionary
    whatweb_output = result_holder.get('whatweb', '')
    nmap_output = result_holder.get('nmap', '')
    nikto_output = result_holder.get('nikto', '')
    dirb_output = result_holder.get('dirb', '')

    # Generate the Word document with the results
    generate_word_document(ip, dir_name, whatweb_output, nmap_output, nikto_output, dirb_output)

if __name__ == "__main__":
    main()
